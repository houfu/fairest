#  Copyright (c) 2021.  Ang Hou Fu
#
#  This software is licensed under the The MIT License.
#  You should have received a copy of the license terms with the software.
#  Otherwise, you can find the text here: https://opensource.org/licenses/MIT
#
#
#  This software is licensed under the The MIT License.
#  You should have received a copy of the license terms with the software.
#  Otherwise, you can find the text here: https://opensource.org/licenses/MIT
#

from io import BytesIO
from typing import Optional, Union, List

from docx import Document
from docx.oxml import CT_P, CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from fairest.models import DocumentModel, DocumentSection, Request, BaseDocumentModelRule, RuleDescription


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    See: https://github.com/python-openxml/python-docx/issues/40
    """
    from docx.document import Document
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)


class BasicDOCXSection(DocumentSection):
    def __init__(self, source: Paragraph, index: str):
        self.source = source
        self.position = index

    def get_text(self) -> str:
        return self.source.text

    def get_position(self) -> str:
        return self.position


class BasicDOCXModel(DocumentModel):

    def __init__(self, source):
        super().__init__()
        self.source = Document(source)
        self.paragraphs = [BasicDOCXSection(block, str(index)) for index, block in
                           enumerate(iter_block_items(self.source))]

    def __getitem__(self, i: int) -> DocumentSection:
        return self.paragraphs[i]

    def __len__(self) -> int:
        return len(self.paragraphs)

    def get_full_text(self) -> List[str]:
        return [paragraph.get_text() for paragraph in self.paragraphs]


class BasicDocxModelRule(BaseDocumentModelRule):

    def check_document(self, document: Union[str, bytes], current: Optional[DocumentModel]) -> bool:
        return isinstance(document, bytes) and current is None

    def run_document_model_rule(self, request: Request) -> Optional[DocumentModel]:
        try:
            return BasicDOCXModel(BytesIO(request.request_body))
        except Exception:
            return None

    @classmethod
    def describe(cls) -> RuleDescription:
        return RuleDescription(
            title="Basic DOCX Model Rule",
            description="The BasicDocX Model reads DOCX (Microsoft Office Word) files"
                        "and transforms it into a model for Fairest processing.",
            author="Core Fairest Plugin"
        )
